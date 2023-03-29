import tkinter as tk
from tkinter import ttk
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches



def generate_memorandum():
    memo_number = memo_number_entry.get()
    to = to_entry.get()
    subject = subject_entry.get()
    driver_name = driver_name_entry.get()
    route = route_entry.get()
    date = date_entry.get()
    departure_time = departure_time_entry.get()
    return_time = return_time_entry.get()
    arrival_time = arrival_time_entry.get()

    doc = Document()

    # Adicionando as imagens
    doc.add_picture('imagem_esquerda.jpg', width=Inches(1.0))
    header_paragraph = doc.add_paragraph()
    header_paragraph.add_run('ESTADO DO PARÁ\nPREFEITURA MUNICIPAL DE CANAÃ DOS CARAJÁS\nSECRETARIA MUNICIPAL DE EDUCAÇÃO\nRua Itamarati, s/n, Novo Horizonte – Canaã dos Carajás – PA, CEP: 68537-000\n\n')
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture('imagem_direita.jpg', width=Inches(1.0))

    # Formatação do corpo do texto
    memo_paragraph = doc.add_paragraph()
    memo_paragraph.add_run(f'MEMORANDO Nº {memo_number} – SECRETARIA DE EDUCAÇÃO\n\nA: {to}\nASSUNTO: {subject}\n\n').bold = True
    memo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    body_paragraph = doc.add_paragraph(f'Prezado(a) Senhor(a),\n\tSolicitamos uma diária para o motorista {driver_name}. Segue abaixo a descrição da viagem:\nRota: {route}\nData: {date}\nHorário de ida: {departure_time}\nHorário de retorno: {return_time}\nHorário de chegada: {arrival_time}\n\nDesde já, agradeço à atenção.')
    body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    signature_paragraph = doc.add_paragraph('Danilo de Morais Gustavo\nGestor de Transporte Escolar\nPortaria n.º 118/2023 – GP')
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f'Memorando_{memo_number}.docx')

app = tk.Tk()
app.title('Gerador de Memorando')

memo_number_label = ttk.Label(app, text='Número do Memorando:')
memo_number_label.grid(column=0, row=0)
memo_number_entry = ttk.Entry(app)
memo_number_entry.grid(column=1, row=0)

to_label = ttk.Label(app, text='Para:')
to_label.grid(column=0, row=1)
to_entry = ttk.Entry(app)
to_entry.grid(column=1, row=1)

subject_label = ttk.Label(app, text='Assunto:')
subject_label.grid(column=0, row=2)
subject_entry = ttk.Entry(app)
subject_entry.grid(column=1, row=2)

driver_name_label = ttk.Label(app, text='Nome do Motorista:')
driver_name_label.grid(column=0, row=3)
driver_name_entry = ttk.Entry(app)
driver_name_entry.grid(column=1, row=3)

route_label = ttk.Label(app, text='Rota:')
route_label.grid(column=0, row=4)
route_entry = ttk.Entry(app)
route_entry.grid(column=1, row=4)

date_label = ttk.Label(app, text='Data:')
date_label.grid(column=0, row=5)
date_entry = ttk.Entry(app)
date_entry.grid(column=1, row=5)

departure_time_label = ttk.Label(app, text='Horário de ida:')
departure_time_label.grid(column=0, row=6)
departure_time_entry = ttk.Entry(app)
departure_time_entry.grid(column=1, row=6)

return_time_label = ttk.Label(app, text='Horário de retorno:')
return_time_label.grid(column=0, row=7)
return_time_entry = ttk.Entry(app)
return_time_entry.grid(column=1, row=7)

arrival_time_label = ttk.Label(app, text='Horário de chegada:')
arrival_time_label.grid(column=0, row=8)
arrival_time_entry = ttk.Entry(app)
arrival_time_entry.grid(column=1, row=8)

generate_button = ttk.Button(app, text='Gerar Memorando', command=generate_memorandum)
generate_button.grid(column=0, row=9, columnspan=2)

app.mainloop()

